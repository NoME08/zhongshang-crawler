"""
从 data.json 重新生成格式化 docx，完全参照 output_clean_final.docx 样式。

格式规范：
  - 页面：A4，上下边距 2.54cm，左右边距 3.18cm
  - 每篇文章独立节，页眉各自显示该篇发文日期
  - 页眉：「来源：转载自中商产业研究院 {发文日期}发文，仅供内部学习使用」（方正仿宋 7.5pt 居中）
  - 文章标题：Arial / 黑体 20pt 加粗 左对齐
  - 子标题：Arial / 黑体 14pt 加粗 左对齐（「一、xxx」「1.xxx」「（1）xxx」）
  - 正文：Arial / 宋体 10.5pt，首行缩进 2 字符
  - 图片：14cm 宽，居中
  - 数据来源：灰色 9pt
"""

import json
import os
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger("reformat")


def is_sub_heading(text):
    """判断是否为子标题"""
    if len(text) > 40:
        return False
    patterns = [
        r'^[一二三四五六七八九十]、',
        r'^\d+[\.、]',
        r'^（[一二三四五六七八九十\d]+）',
        r'^\([一二三四五六七八九十\d]+\)',
        r'^[（\(][一二三四五六七八九十]+[）\)]',
    ]
    for pat in patterns:
        if re.match(pat, text):
            return True
    return False


def is_data_source(text):
    return "数据来源" in text or "资料来源" in text


def set_run_font(run, western, east_asian, size, bold=False, color=None):
    """统一设置 run 的字体"""
    run.font.name = western
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asian)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_page_number(paragraph):
    """在段落中添加分节页码字段（WORD PAGE field，每节独立计数）"""
    from docx.oxml import OxmlElement

    run1 = paragraph.add_run()
    run1.font.size = Pt(9)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fld_begin)

    run2 = paragraph.add_run()
    run2.font.size = Pt(9)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    run3.font.size = Pt(9)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run3._element.append(fld_end)


def _setup_section(section):
    """设置节的页面边距"""
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def _set_header(section, publish_time):
    """设置节的页眉，使用文章的发文日期"""
    date_str = ""
    if publish_time:
        try:
            dt = datetime.strptime(publish_time[:10], "%Y-%m-%d")
            date_str = f"{dt.year}年{dt.month}月"
        except Exception:
            pass
    if not date_str:
        now = datetime.now()
        date_str = f"{now.year}年{now.month}月"

    header = section.header
    header.is_linked_to_previous = False
    for para in header.paragraphs:
        para.clear()
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(
        f"来源：转载自中商产业研究院 {date_str}发文，仅供内部学习使用"
    )
    set_run_font(run, "方正仿宋_GB2312", "方正仿宋_GB2312", Pt(7.5))

    # ---- 页脚：分节页码 ----
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)


def generate_formatted_docx(data_path, image_dir, output_path):
    """从 data.json 生成格式化的 Word 文档"""
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    articles = data.get("articles", [])
    total = len(articles)

    for idx, article in enumerate(articles):
        # ---- 每篇文章新节（新页 + 独立页眉） ----
        if idx > 0:
            doc.add_section()
        section = doc.sections[-1]
        _setup_section(section)
        _set_header(section, article.get("publish_time", ""))

        # ---- 文章标题 ----
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(6)
        run = title_para.add_run(article.get("title", ""))
        set_run_font(run, "Arial", "黑体", Pt(20), bold=True)

        # ---- 正文内容 ----
        for item in article.get("content", []):
            if item["type"] == "text":
                text = item["value"].strip()
                if not text:
                    continue

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if is_sub_heading(text):
                    para.paragraph_format.space_before = Pt(8)
                    para.paragraph_format.space_after = Pt(4)
                    run = para.add_run(text)
                    set_run_font(run, "Arial", "黑体", Pt(14), bold=True)
                elif is_data_source(text):
                    run = para.add_run(text)
                    set_run_font(run, "Arial", "宋体", Pt(9),
                                 color=RGBColor(0x80, 0x80, 0x80))
                else:
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    run = para.add_run(text)
                    set_run_font(run, "Arial", "宋体", Pt(10.5))

            elif item["type"] == "image":
                local_path = item.get("local_path", "")
                if not local_path or not os.path.exists(local_path):
                    continue

                try:
                    from PIL import Image as PILImage
                    img = PILImage.open(local_path)
                    w, h = img.size
                    if h > 0 and w / h > 3.5:
                        continue
                except Exception:
                    pass

                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(local_path, width=Cm(14))
                except Exception as e:
                    logger.warning(f"插入图片失败: {local_path} - {e}")

        # ---- 数据来源（文章末尾） ----
        data_source = article.get("data_source", "")
        if data_source:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(data_source)
            set_run_font(run, "Arial", "宋体", Pt(9),
                         color=RGBColor(0x80, 0x80, 0x80))

    doc.save(output_path)
    print(f"格式化完成: {output_path} ({total} 篇文章)")


def _check_libreoffice():
    """检测 LibreOffice 是否可用，返回 soffice 路径或 None

    不仅检查文件是否存在，还实际运行 --version 做功能验证。
    避免"文件在但被 macOS 隔离/权限不足/配置损坏"导致的误判。
    """
    import subprocess
    import shutil

    # 收集所有候选路径
    candidates = []

    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        candidates.append(mac_path)

    for win_base in ["C:\\Program Files\\LibreOffice\\program",
                     "C:\\Program Files (x86)\\LibreOffice\\program"]:
        win_path = os.path.join(win_base, "soffice.exe")
        if os.path.exists(win_path):
            candidates.append(win_path)

    if shutil.which("soffice"):
        candidates.append("soffice")
    if shutil.which("libreoffice"):
        candidates.append("libreoffice")

    # 对每个候选做功能验证：实际运行 --headless --version
    for candidate in candidates:
        try:
            result = subprocess.run(
                [candidate, "--headless", "--version"],
                timeout=15,
                capture_output=True,
                text=True,
            )
            if result.returncode == 0 and "LibreOffice" in result.stdout:
                logger.debug(f"LibreOffice 功能验证通过: {candidate}")
                return candidate
            else:
                logger.debug(
                    f"LibreOffice 候选 {candidate} 验证失败 "
                    f"(rc={result.returncode}, stderr={result.stderr[:100]})"
                )
        except FileNotFoundError:
            logger.debug(f"LibreOffice 候选 {candidate} 文件不存在")
        except subprocess.TimeoutExpired:
            logger.debug(f"LibreOffice 候选 {candidate} 验证超时（可能被 macOS 隔离）")
        except Exception as e:
            logger.debug(f"LibreOffice 候选 {candidate} 验证异常: {e}")

    return None


def export_pdfs(data_path, image_dir, output_dir):
    """每篇文章导出为独立 PDF（通过 LibreOffice 转换）"""
    import subprocess
    import time

    soffice = _check_libreoffice()
    if not soffice:
        raise RuntimeError(
            "未找到可用的 LibreOffice，无法转换为 PDF。\n"
            "请安装 LibreOffice：https://www.libreoffice.org/download/\n"
            "  macOS: brew install --cask libreoffice\n"
            "  Windows: 从官网下载安装包\n"
            "\n"
            "如果已安装但仍报错（常见于 macOS）：\n"
            "  1. 尝试手动打开一次 LibreOffice.app（Finder → 应用程序 → LibreOffice）\n"
            "  2. 或在终端执行：xattr -d com.apple.quarantine /Applications/LibreOffice.app\n"
            "  3. 然后重新运行本程序"
        )

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    os.makedirs(output_dir, exist_ok=True)
    articles = data.get("articles", [])
    total = len(articles)
    success = 0

    for idx, article in enumerate(articles):
        title = article.get("title", f"article_{idx}")
        safe = re.sub(r'[\\/:*?"<>|]', '_', title)[:80]
        docx_path = os.path.join(output_dir, f"{safe}.docx")
        pdf_path = os.path.join(output_dir, f"{safe}.pdf")

        # 生成单篇 docx
        try:
            _write_single_docx(article, image_dir, docx_path)
        except Exception as e:
            logger.warning(f"  [{idx+1}/{total}] 生成 docx 失败: {e}")
            continue

        # 转 PDF（含重试）
        pdf_ok = False
        for attempt in range(3):
            try:
                result = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", output_dir, docx_path],
                    timeout=90, capture_output=True, text=True,
                )
                # LibreOffice 可能返回 0 但仍然转换失败
                # 必须验证 PDF 文件确实被创建
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    pdf_ok = True
                    break
                else:
                    stderr = result.stderr.strip()
                    stdout = result.stdout.strip()
                    logger.warning(
                        f"  [{idx+1}/{total}] LibreOffice 返回码 {result.returncode} "
                        f"但未生成 PDF (attempt {attempt+1})"
                    )
                    if stderr:
                        logger.warning(f"    stderr: {stderr[:200]}")
                    if stdout:
                        logger.warning(f"    stdout: {stdout[:200]}")
            except subprocess.TimeoutExpired:
                logger.warning(
                    f"  [{idx+1}/{total}] 转换超时 90s (attempt {attempt+1})"
                )
            except FileNotFoundError:
                raise RuntimeError(
                    f"找不到 LibreOffice 可执行文件: {soffice}\n"
                    "请确认 LibreOffice 已正确安装"
                )
            except Exception as e:
                logger.warning(
                    f"  [{idx+1}/{total}] 转换异常 (attempt {attempt+1}): {e}"
                )
            if attempt < 2:
                time.sleep(2)

        if pdf_ok:
            os.remove(docx_path)
            success += 1
        else:
            # 转换失败：删除 docx，避免被打包进 zip
            logger.error(
                f"  [{idx+1}/{total}] PDF 转换失败（已重试 3 次），跳过: {safe[:50]}"
            )
            try:
                os.remove(docx_path)
            except Exception:
                pass

    print(f"PDF 已输出到 {output_dir}/ ({success}/{total} 篇，{total - success} 篇失败)")


def _write_single_docx(article, image_dir, output_path):
    """生成单篇文章的完整格式化 docx"""
    doc = Document()
    section = doc.sections[0]
    _setup_section(section)
    _set_header(section, article.get("publish_time", ""))

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run(article.get("title", ""))
    set_run_font(run, "Arial", "黑体", Pt(20), bold=True)

    # 正文
    for item in article.get("content", []):
        if item["type"] == "text":
            text = item["value"].strip()
            if not text:
                continue
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if is_sub_heading(text):
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run(text)
                set_run_font(run, "Arial", "黑体", Pt(14), bold=True)
            elif is_data_source(text):
                run = para.add_run(text)
                set_run_font(run, "Arial", "宋体", Pt(9), color=RGBColor(0x80, 0x80, 0x80))
            else:
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.first_line_indent = Cm(0.74)
                run = para.add_run(text)
                set_run_font(run, "Arial", "宋体", Pt(10.5))
        elif item["type"] == "image":
            local_path = item.get("local_path", "")
            if not local_path or not os.path.exists(local_path):
                continue
            try:
                from PIL import Image as PILImage
                img = PILImage.open(local_path)
                w, h = img.size
                if h > 0 and w / h > 3.5:
                    continue
            except Exception:
                pass
            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(local_path, width=Cm(14))
            except Exception:
                pass

    # 数据来源
    ds = article.get("data_source", "")
    if ds:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(ds)
        set_run_font(run, "Arial", "宋体", Pt(9), color=RGBColor(0x80, 0x80, 0x80))

    doc.save(output_path)


def generate_titles_docx(data_path, output_path):
    """从 data.json 生成仅包含文章标题的 Word 文档"""
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    section = doc.sections[0]
    _setup_section(section)

    # 大标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    today = datetime.now().strftime("%Y年%m月%d日")
    run = title_para.add_run(f"中商情报网 · 文章标题汇总（{today}）")
    set_run_font(run, "Arial", "黑体", Pt(18), bold=True)

    articles = data.get("articles", [])
    for idx, article in enumerate(articles, 1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(f"{idx}. {article.get('title', '')}")
        set_run_font(run, "Arial", "宋体", Pt(12))

    doc.save(output_path)
    print(f"标题文档已生成: {output_path} ({len(articles)} 条标题)")


if __name__ == "__main__":
    import sys
    data_path = sys.argv[1] if len(sys.argv) > 1 else "output/data.json"
    image_dir = sys.argv[2] if len(sys.argv) > 2 else "output/images"
    output_path = sys.argv[3] if len(sys.argv) > 3 else "output/output_formatted.docx"
    generate_formatted_docx(data_path, image_dir, output_path)
